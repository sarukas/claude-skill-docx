#!/usr/bin/env python3
"""
Find and replace text in DOCX documents.

Supports simple replacement (preserving formatting) and tracked-changes mode.

Usage:
    python docx_find_replace.py input.docx output.docx --find "Old" --replace "New"
    python docx_find_replace.py input.docx output.docx --find "Old" --replace "New" --track-changes
    python docx_find_replace.py input.docx output.docx --find "Old" --replace "New" --dry-run
    python docx_find_replace.py input.docx output.docx --find "Old" --replace "New" --scope tables
"""

import argparse
import re
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Missing dependency: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    import defusedxml.minidom
except ImportError:
    print("Missing dependency: pip install defusedxml", file=sys.stderr)
    sys.exit(1)


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _replace_in_runs(paragraph, find_text, replace_text, case_sensitive=True, whole_word=False):
    """Replace text across runs in a paragraph, preserving formatting.

    Returns the number of replacements made.
    """
    # Build full paragraph text from runs
    runs = paragraph.runs
    if not runs:
        return 0

    full_text = "".join(r.text for r in runs)

    # Build regex pattern
    flags = 0 if case_sensitive else re.IGNORECASE
    pattern = re.escape(find_text)
    if whole_word:
        pattern = r"\b" + pattern + r"\b"

    matches = list(re.finditer(pattern, full_text, flags))
    if not matches:
        return 0

    # Map character positions to (run_index, char_within_run)
    char_map = []
    for r_idx, run in enumerate(runs):
        for c_idx in range(len(run.text)):
            char_map.append((r_idx, c_idx))

    # Replace from last to first to preserve positions
    for match in reversed(matches):
        start, end = match.start(), match.end()

        # Find which runs are affected
        start_run, start_char = char_map[start]
        end_run, end_char = char_map[end - 1]

        if start_run == end_run:
            # Simple case: match is within a single run
            run = runs[start_run]
            run.text = run.text[:start_char] + replace_text + run.text[end_char + 1:]
        else:
            # Complex case: match spans multiple runs
            # Put replacement text in the first run, clear the rest
            runs[start_run].text = runs[start_run].text[:start_char] + replace_text
            runs[end_run].text = runs[end_run].text[end_char + 1:]
            for r_idx in range(start_run + 1, end_run):
                runs[r_idx].text = ""

        # Rebuild char_map after each replacement
        char_map = []
        for r_idx, run in enumerate(runs):
            for c_idx in range(len(run.text)):
                char_map.append((r_idx, c_idx))

    return len(matches)


def simple_replace(input_path, output_path, find_text, replace_text,
                   case_sensitive=True, whole_word=False, scope="all", dry_run=False):
    """Replace text using python-docx, preserving formatting."""
    doc = Document(str(input_path))
    total = 0

    if scope in ("all", "body"):
        for para in doc.paragraphs:
            count = _replace_in_runs(para, find_text, replace_text, case_sensitive, whole_word)
            total += count

    if scope in ("all", "tables"):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        count = _replace_in_runs(para, find_text, replace_text, case_sensitive, whole_word)
                        total += count

    if scope in ("all", "headers", "footers"):
        for section in doc.sections:
            if scope in ("all", "headers"):
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header and header.is_linked_to_previous is False:
                        for para in header.paragraphs:
                            count = _replace_in_runs(para, find_text, replace_text, case_sensitive, whole_word)
                            total += count

            if scope in ("all", "footers"):
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer and footer.is_linked_to_previous is False:
                        for para in footer.paragraphs:
                            count = _replace_in_runs(para, find_text, replace_text, case_sensitive, whole_word)
                            total += count

    if dry_run:
        print(f"Dry run: {total} occurrences of \"{find_text}\" found")
    else:
        doc.save(str(output_path))
        print(f"Replaced {total} occurrences of \"{find_text}\" with \"{replace_text}\"")
        print(f"Saved to: {output_path}")

    return total


def tracked_replace(input_path, output_path, find_text, replace_text,
                    case_sensitive=True, whole_word=False, author="RIO AI", dry_run=False):
    """Replace text with tracked changes using OOXML manipulation.

    Uses the unpack -> edit -> pack workflow via the Document class.
    """
    import random
    from datetime import datetime, timezone

    # Add parent dir to path for imports
    script_dir = Path(__file__).resolve().parent
    skill_dir = script_dir.parent
    sys.path.insert(0, str(skill_dir))
    sys.path.insert(0, str(script_dir))

    from ooxml.scripts.pack import pack_document

    # Unpack to temp directory
    temp_dir = tempfile.mkdtemp()
    unpacked = Path(temp_dir) / "unpacked"
    unpacked.mkdir()

    with zipfile.ZipFile(input_path, "r") as zf:
        zf.extractall(unpacked)

    doc_xml_path = unpacked / "word" / "document.xml"
    if not doc_xml_path.exists():
        print("Error: word/document.xml not found in DOCX", file=sys.stderr)
        shutil.rmtree(temp_dir)
        sys.exit(1)

    # Parse document XML
    dom = defusedxml.minidom.parse(str(doc_xml_path))
    rsid = "".join(random.choices("0123456789ABCDEF", k=8))
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    # Find all w:t elements and replace
    total = 0
    flags = 0 if case_sensitive else re.IGNORECASE
    pattern = re.escape(find_text)
    if whole_word:
        pattern = r"\b" + pattern + r"\b"

    change_id = 100  # Start high to avoid collisions

    for t_elem in list(dom.getElementsByTagName("w:t")):
        text = t_elem.firstChild.nodeValue if t_elem.firstChild else ""
        if not text:
            continue

        matches = list(re.finditer(pattern, text, flags))
        if not matches:
            continue

        if dry_run:
            total += len(matches)
            continue

        # Get the parent w:r element
        r_elem = t_elem.parentNode
        p_elem = r_elem.parentNode

        # For each match, split the text and create tracked changes
        # Process from last to first to preserve positions
        for match in reversed(matches):
            total += 1
            start, end = match.start(), match.end()

            before = text[:start]
            matched = text[start:end]
            after = text[end:]

            # Get run properties (rPr) from original run
            rpr_nodes = r_elem.getElementsByTagName("w:rPr")
            rpr_xml = ""
            if rpr_nodes:
                rpr_xml = rpr_nodes[0].toxml()

            # Build replacement XML fragments
            fragments = []

            if before:
                fragments.append(f'<w:r w:rsidR="{rsid}">{rpr_xml}<w:t xml:space="preserve">{before}</w:t></w:r>')

            # Deletion
            fragments.append(
                f'<w:del w:id="{change_id}" w:author="{author}" w:date="{now}">'
                f'<w:r w:rsidDel="{rsid}">{rpr_xml}<w:delText xml:space="preserve">{matched}</w:delText></w:r>'
                f'</w:del>'
            )
            change_id += 1

            # Insertion
            fragments.append(
                f'<w:ins w:id="{change_id}" w:author="{author}" w:date="{now}">'
                f'<w:r w:rsidR="{rsid}">{rpr_xml}<w:t xml:space="preserve">{replace_text}</w:t></w:r>'
                f'</w:ins>'
            )
            change_id += 1

            if after:
                fragments.append(f'<w:r w:rsidR="{rsid}">{rpr_xml}<w:t xml:space="preserve">{after}</w:t></w:r>')

            # Parse fragments and replace the original run
            ns_decl = []
            root = dom.documentElement
            if root.attributes:
                for i in range(root.attributes.length):
                    attr = root.attributes.item(i)
                    if attr.name.startswith("xmlns"):
                        ns_decl.append(f'{attr.name}="{attr.value}"')
            ns_str = " ".join(ns_decl)

            wrapper_xml = f'<root {ns_str}>{"".join(fragments)}</root>'
            frag_doc = defusedxml.minidom.parseString(wrapper_xml)
            new_nodes = [dom.importNode(child, True) for child in frag_doc.documentElement.childNodes]

            for node in new_nodes:
                p_elem.insertBefore(node, r_elem)
            p_elem.removeChild(r_elem)

            # Update text for next iteration (already processing reversed)
            text = before if before else ""

    if dry_run:
        print(f"Dry run: {total} occurrences of \"{find_text}\" found")
        shutil.rmtree(temp_dir)
        return total

    # Save modified XML
    doc_xml_path.write_bytes(dom.toxml(encoding="utf-8"))

    # Pack back to DOCX
    pack_document(str(unpacked), str(output_path))
    print(f"Replaced {total} occurrences with tracked changes (author: {author})")
    print(f"Saved to: {output_path}")

    shutil.rmtree(temp_dir)
    return total


def main():
    parser = argparse.ArgumentParser(
        description="Find and replace text in DOCX documents",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("input", type=Path, help="Input DOCX file")
    parser.add_argument("output", type=Path, help="Output DOCX file")
    parser.add_argument("--find", required=True, help="Text to find")
    parser.add_argument("--replace", required=True, help="Replacement text")
    parser.add_argument("--case-sensitive", action="store_true", default=True,
                        help="Case-sensitive search (default)")
    parser.add_argument("--no-case-sensitive", dest="case_sensitive", action="store_false",
                        help="Case-insensitive search")
    parser.add_argument("--whole-word", action="store_true", help="Match whole words only")
    parser.add_argument("--track-changes", action="store_true",
                        help="Use tracked changes instead of direct replacement")
    parser.add_argument("--author", default="RIO AI",
                        help="Author name for tracked changes (default: RIO AI)")
    parser.add_argument("--dry-run", action="store_true",
                        help="Report matches without modifying")
    parser.add_argument("--scope", choices=["body", "headers", "footers", "tables", "all"],
                        default="all", help="Where to search (default: all)")

    args = parser.parse_args()

    if not args.input.is_file():
        print(f"Error: File not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    if args.track_changes:
        tracked_replace(
            args.input, args.output, args.find, args.replace,
            case_sensitive=args.case_sensitive, whole_word=args.whole_word,
            author=args.author, dry_run=args.dry_run,
        )
    else:
        simple_replace(
            args.input, args.output, args.find, args.replace,
            case_sensitive=args.case_sensitive, whole_word=args.whole_word,
            scope=args.scope, dry_run=args.dry_run,
        )


if __name__ == "__main__":
    main()
