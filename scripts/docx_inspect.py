#!/usr/bin/env python3
"""
Inspect DOCX document structure and content.

Outputs Markdown to stdout for easy analysis by Claude.

Usage:
    python docx_inspect.py input.docx                    # Structure summary (default)
    python docx_inspect.py input.docx --text              # Paragraph text with indices
    python docx_inspect.py input.docx --headings          # Heading outline
    python docx_inspect.py input.docx --tables            # Tables as Markdown
    python docx_inspect.py input.docx --comments          # Comments with metadata
    python docx_inspect.py input.docx --tracked-changes   # Insertions/deletions
    python docx_inspect.py input.docx --structure         # Document stats
    python docx_inspect.py input.docx --text --headings   # Multiple modes
"""

import argparse
import sys
import zipfile
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Missing dependency: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    import defusedxml.minidom
except ImportError:
    defusedxml = None


# OOXML namespaces
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _get_text(elem):
    """Recursively extract text from an XML element (lxml or minidom)."""
    parts = []
    # Detect whether this is a minidom node or an lxml element
    if hasattr(elem, 'childNodes'):
        # minidom DOM node
        for child in elem.childNodes:
            tag = getattr(child, 'nodeName', '') or ''
            if tag == 'w:t':
                parts.append(child.firstChild.nodeValue if child.firstChild else "")
            elif tag == 'w:tab':
                parts.append("\t")
            elif tag == 'w:br':
                parts.append("\n")
            elif tag == 'w:r':
                parts.append(_get_text(child))
            elif tag == 'w:delText':
                parts.append(child.firstChild.nodeValue if child.firstChild else "")
    else:
        # lxml element
        for child in elem:
            if child.tag == f"{{{W}}}t":
                parts.append(child.text or "")
            elif child.tag == f"{{{W}}}tab":
                parts.append("\t")
            elif child.tag == f"{{{W}}}br":
                parts.append("\n")
            elif child.tag == f"{{{W}}}r":
                parts.append(_get_text(child))
            elif child.tag == f"{{{W}}}delText":
                parts.append(child.text or "")
    return "".join(parts)


def inspect_text(doc):
    """Print paragraph-by-paragraph text with indices."""
    print("## Document Text\n")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else "None"
        if text:
            print(f"**[{i}]** ({style}) {text}\n")
        else:
            print(f"**[{i}]** ({style}) _(empty)_\n")


def inspect_headings(doc):
    """Print heading outline with hierarchy."""
    print("## Heading Outline\n")
    found = False
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        if style.startswith("Heading"):
            found = True
            try:
                level = int(style.split()[-1])
            except (ValueError, IndexError):
                level = 1
            indent = "  " * (level - 1)
            text = para.text.strip() or "_(empty heading)_"
            print(f"{indent}- **H{level}** [{i}]: {text}")
    if not found:
        print("_No headings found in document._")
    print()


def inspect_tables(doc):
    """Print tables as Markdown with row/column indices."""
    print("## Tables\n")
    if not doc.tables:
        print("_No tables found in document._\n")
        return

    for t_idx, table in enumerate(doc.tables):
        print(f"### Table {t_idx + 1} ({len(table.rows)} rows x {len(table.columns)} columns)\n")
        for r_idx, row in enumerate(table.rows):
            cells = []
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip().replace("\n", " ")
                cells.append(text)
            print(f"| {' | '.join(cells)} |")
            if r_idx == 0:
                print(f"| {' | '.join(['---'] * len(cells))} |")
        print()


def inspect_comments(docx_path):
    """Extract and print comments from DOCX."""
    print("## Comments\n")

    try:
        zf = zipfile.ZipFile(docx_path, "r")
    except zipfile.BadZipFile:
        print("_Error: Not a valid DOCX file._\n")
        return

    # Parse comments.xml
    if "word/comments.xml" not in zf.namelist():
        print("_No comments found in document._\n")
        zf.close()
        return

    comments_xml = zf.read("word/comments.xml")
    if defusedxml:
        dom = defusedxml.minidom.parseString(comments_xml)
    else:
        from xml.dom import minidom
        dom = minidom.parseString(comments_xml)

    comments = dom.getElementsByTagName("w:comment")

    # Parse commentsExtended.xml for resolved state
    resolved_map = {}
    if "word/commentsExtended.xml" in zf.namelist():
        ext_xml = zf.read("word/commentsExtended.xml")
        if defusedxml:
            ext_dom = defusedxml.minidom.parseString(ext_xml)
        else:
            from xml.dom import minidom
            ext_dom = minidom.parseString(ext_xml)
        for ex in ext_dom.getElementsByTagName("w15:commentEx"):
            para_id = ex.getAttribute("w15:paraId")
            done = ex.getAttribute("w15:done")
            if para_id:
                resolved_map[para_id] = done == "1"

    # Parse document.xml to find anchor text for each comment
    anchor_map = {}
    if "word/document.xml" in zf.namelist():
        doc_xml = zf.read("word/document.xml")
        if defusedxml:
            doc_dom = defusedxml.minidom.parseString(doc_xml)
        else:
            from xml.dom import minidom
            doc_dom = minidom.parseString(doc_xml)

        # Find commentRangeStart/End pairs
        range_starts = {}
        for rs in doc_dom.getElementsByTagName("w:commentRangeStart"):
            cid = rs.getAttribute("w:id")
            if cid:
                range_starts[cid] = rs

        for cid, rs in range_starts.items():
            # Collect text between rangeStart and rangeEnd
            texts = []
            sibling = rs.nextSibling
            while sibling is not None:
                if sibling.nodeType == sibling.ELEMENT_NODE:
                    if sibling.tagName == "w:commentRangeEnd" and sibling.getAttribute("w:id") == cid:
                        break
                    if sibling.tagName == "w:r":
                        for t in sibling.getElementsByTagName("w:t"):
                            if t.firstChild:
                                texts.append(t.firstChild.nodeValue or "")
                sibling = sibling.nextSibling
            if texts:
                anchor_map[cid] = "".join(texts)

    if not comments:
        print("_No comments found._\n")
        zf.close()
        return

    for comment in comments:
        cid = comment.getAttribute("w:id")
        author = comment.getAttribute("w:author") or "Unknown"
        date = comment.getAttribute("w:date") or "Unknown"

        # Get comment text
        text_parts = []
        for t in comment.getElementsByTagName("w:t"):
            if t.firstChild:
                text_parts.append(t.firstChild.nodeValue or "")
        text = "".join(text_parts)

        # Check resolved state via paraId
        resolved = False
        for p in comment.getElementsByTagName("w:p"):
            para_id = p.getAttribute("w14:paraId")
            if para_id and para_id in resolved_map:
                resolved = resolved_map[para_id]
                break

        anchor = anchor_map.get(cid, "_(anchor not found)_")
        status = "RESOLVED" if resolved else "OPEN"

        print(f"**Comment {cid}** [{status}]")
        print(f"- Author: {author}")
        print(f"- Date: {date}")
        print(f"- Anchor: \"{anchor}\"")
        print(f"- Text: {text}")
        print()

    zf.close()


def inspect_tracked_changes(docx_path):
    """Extract and print tracked changes."""
    print("## Tracked Changes\n")

    try:
        zf = zipfile.ZipFile(docx_path, "r")
    except zipfile.BadZipFile:
        print("_Error: Not a valid DOCX file._\n")
        return

    if "word/document.xml" not in zf.namelist():
        print("_No document.xml found._\n")
        zf.close()
        return

    doc_xml = zf.read("word/document.xml")
    if defusedxml:
        dom = defusedxml.minidom.parseString(doc_xml)
    else:
        from xml.dom import minidom
        dom = minidom.parseString(doc_xml)

    insertions = dom.getElementsByTagName("w:ins")
    deletions = dom.getElementsByTagName("w:del")

    if not insertions and not deletions:
        print("_No tracked changes found._\n")
        zf.close()
        return

    if insertions:
        print(f"### Insertions ({len(insertions)})\n")
        for ins in insertions:
            cid = ins.getAttribute("w:id") or "?"
            author = ins.getAttribute("w:author") or "Unknown"
            date = ins.getAttribute("w:date") or "Unknown"
            text = _get_text(ins)
            if text.strip():
                print(f"- **[{cid}]** by {author} ({date}): ++{text.strip()}++")
        print()

    if deletions:
        print(f"### Deletions ({len(deletions)})\n")
        for del_elem in deletions:
            cid = del_elem.getAttribute("w:id") or "?"
            author = del_elem.getAttribute("w:author") or "Unknown"
            date = del_elem.getAttribute("w:date") or "Unknown"
            # Get deleted text from w:delText
            text_parts = []
            for dt in del_elem.getElementsByTagName("w:delText"):
                if dt.firstChild:
                    text_parts.append(dt.firstChild.nodeValue or "")
            text = "".join(text_parts)
            if text.strip():
                print(f"- **[{cid}]** by {author} ({date}): ~~{text.strip()}~~")
        print()

    zf.close()


def inspect_structure(doc, docx_path):
    """Print document structure summary."""
    print("## Document Structure\n")

    # Paragraph count
    total_paras = len(doc.paragraphs)
    non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
    print(f"- **Paragraphs**: {total_paras} total ({non_empty} non-empty)")

    # Table count
    print(f"- **Tables**: {len(doc.tables)}")

    # Section count
    print(f"- **Sections**: {len(doc.sections)}")

    # Style inventory
    styles_used = {}
    for para in doc.paragraphs:
        name = para.style.name if para.style else "None"
        styles_used[name] = styles_used.get(name, 0) + 1

    print(f"\n### Style Inventory ({len(styles_used)} styles)\n")
    for style, count in sorted(styles_used.items(), key=lambda x: -x[1]):
        print(f"- {style}: {count}")

    # Check for special content
    print("\n### Special Content\n")

    try:
        zf = zipfile.ZipFile(docx_path, "r")
        names = zf.namelist()
        has_comments = "word/comments.xml" in names
        has_headers = any("header" in n for n in names)
        has_footers = any("footer" in n for n in names)
        has_images = any(n.startswith("word/media/") for n in names)
        image_count = sum(1 for n in names if n.startswith("word/media/"))

        print(f"- Comments: {'Yes' if has_comments else 'No'}")
        print(f"- Headers: {'Yes' if has_headers else 'No'}")
        print(f"- Footers: {'Yes' if has_footers else 'No'}")
        print(f"- Embedded images: {image_count}")

        # Check for tracked changes
        if "word/document.xml" in names:
            doc_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
            ins_count = doc_xml.count("<w:ins ")
            del_count = doc_xml.count("<w:del ")
            if ins_count or del_count:
                print(f"- Tracked changes: {ins_count} insertions, {del_count} deletions")
            else:
                print("- Tracked changes: No")

        zf.close()
    except Exception as e:
        print(f"- _(Could not inspect ZIP: {e})_")

    print()


def main():
    parser = argparse.ArgumentParser(
        description="Inspect DOCX document structure and content",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("input", type=Path, help="Input DOCX file")
    parser.add_argument("--text", action="store_true", help="Paragraph text with indices")
    parser.add_argument("--headings", action="store_true", help="Heading outline")
    parser.add_argument("--tables", action="store_true", help="Tables as Markdown")
    parser.add_argument("--comments", action="store_true", help="Comments with metadata")
    parser.add_argument("--tracked-changes", action="store_true", help="Insertions/deletions")
    parser.add_argument("--structure", action="store_true", help="Document stats")

    args = parser.parse_args()

    if not args.input.is_file():
        print(f"Error: File not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    # Default to --structure if no flags given
    any_flag = args.text or args.headings or args.tables or args.comments or args.tracked_changes or args.structure
    if not any_flag:
        args.structure = True

    # Open with python-docx for text/headings/tables/structure
    doc = None
    if args.text or args.headings or args.tables or args.structure:
        try:
            doc = Document(str(args.input))
        except Exception as e:
            print(f"Error opening document: {e}", file=sys.stderr)
            sys.exit(1)

    print(f"# Inspection: {args.input.name}\n")

    if args.structure:
        inspect_structure(doc, args.input)
    if args.headings:
        inspect_headings(doc)
    if args.text:
        inspect_text(doc)
    if args.tables:
        inspect_tables(doc)
    if args.comments:
        inspect_comments(args.input)
    if args.tracked_changes:
        inspect_tracked_changes(args.input)


if __name__ == "__main__":
    main()
