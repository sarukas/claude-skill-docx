#!/usr/bin/env python3
"""
Add comments to a DOCX document from a JSON manifest.

Comments are anchored by matching text in paragraphs. Uses python-docx for
initial processing and ZIP-level patching for comment XML parts.

Usage:
    python docx_add_comments.py input.docx output.docx --comments comments.json
    python docx_add_comments.py input.docx output.docx --comments comments.json --author "RIO AI"

JSON format:
    [
        {"anchor_text": "text to comment on", "text": "Comment content"},
        {"anchor_text": "other text", "text": "Another comment", "resolved": true},
        {"anchor_text": "reply target", "text": "Reply to comment", "reply_to": 0}
    ]

Fields:
    anchor_text: Text in document to attach the comment to (required)
    text: Comment content (required)
    resolved: Whether comment is resolved (optional, default false)
    reply_to: Index of parent comment to reply to (optional, 0-based)
"""

import argparse
import json
import random
import shutil
import sys
import tempfile
import zipfile
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from xml.etree import ElementTree as ET

try:
    from docx import Document
    from lxml import etree as lxml_etree
except ImportError:
    print("Missing dependency: pip install python-docx lxml", file=sys.stderr)
    sys.exit(1)


# OOXML namespaces
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"
CT = "http://schemas.openxmlformats.org/package/2006/content-types"
REL = "http://schemas.openxmlformats.org/package/2006/relationships"

# Register namespaces to avoid ns0/ns1 prefixes in output
ET.register_namespace("w", W)
ET.register_namespace("r", R)
ET.register_namespace("w14", W14)
ET.register_namespace("w15", W15)
ET.register_namespace("mc", MC)


def _gen_para_id():
    """Generate a random 8-character hex paraId."""
    return "".join(random.choices("0123456789ABCDEF", k=8))


def add_comments(input_path, output_path, comments_json, author="RIO AI", initials="RI"):
    """Add comments to a DOCX file.

    Pipeline:
    1. Open with python-docx, inject comment markers into paragraphs
    2. Save with python-docx (preserves styles, headers, images)
    3. Patch ZIP to add comments.xml, commentsExtended.xml, relationships, content types
    """
    # Load comments manifest
    with open(comments_json, "r", encoding="utf-8") as f:
        comments = json.load(f)

    if not comments:
        print("No comments in manifest.")
        return

    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    base_id = 100  # Start high to avoid collisions

    # Open document
    doc = Document(str(input_path))

    # Phase 1: Inject comment markers into paragraphs
    comment_data = []  # Collect data for XML generation
    placed = 0

    for idx, comment in enumerate(comments):
        anchor_text = comment.get("anchor_text", "")
        text = comment.get("text", "")
        resolved = comment.get("resolved", False)
        reply_to = comment.get("reply_to")

        if not anchor_text or not text:
            print(f"Warning: Skipping comment {idx} (missing anchor_text or text)")
            continue

        cid = base_id + idx
        para_id = _gen_para_id()

        # Find paragraph containing anchor text
        found = False
        for para in doc.paragraphs:
            if anchor_text in para.text:
                _inject_markers(para, cid)
                found = True
                break

        # Also search tables
        if not found:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if anchor_text in para.text:
                                _inject_markers(para, cid)
                                found = True
                                break
                        if found:
                            break
                    if found:
                        break
                if found:
                    break

        if not found:
            print(f"Warning: Anchor text not found for comment {idx}: \"{anchor_text[:50]}...\"")
            continue

        comment_data.append({
            "id": cid,
            "text": text,
            "author": author,
            "initials": initials,
            "date": now,
            "para_id": para_id,
            "resolved": resolved,
            "reply_to": reply_to,
        })
        placed += 1

    if placed == 0:
        print("No comments could be placed (no anchor text matches found).")
        return

    # Save with python-docx (preserves all template parts)
    temp_docx = Path(tempfile.mktemp(suffix=".docx"))
    doc.save(str(temp_docx))

    # Phase 2: Patch ZIP to add comment XML parts
    _patch_zip_with_comments(temp_docx, output_path, comment_data)
    temp_docx.unlink()

    print(f"Added {placed} comments (author: {author})")
    print(f"Saved to: {output_path}")


def _inject_markers(para, comment_id):
    """Inject commentRangeStart, commentRangeEnd, and commentReference into a paragraph.

    Wraps all runs in the paragraph with the comment range.
    """
    elem = para._element

    # Find all existing run elements
    runs = list(elem.findall(f"{{{W}}}r"))
    ppr = elem.find(f"{{{W}}}pPr")

    if not runs:
        return

    # Create marker elements (use lxml SubElement since para._element is lxml)
    range_start = lxml_etree.SubElement(elem, f"{{{W}}}commentRangeStart")
    range_start.set(f"{{{W}}}id", str(comment_id))

    range_end = lxml_etree.SubElement(elem, f"{{{W}}}commentRangeEnd")
    range_end.set(f"{{{W}}}id", str(comment_id))

    ref_run = lxml_etree.SubElement(elem, f"{{{W}}}r")
    ref_rpr = lxml_etree.SubElement(ref_run, f"{{{W}}}rPr")
    ref_style = lxml_etree.SubElement(ref_rpr, f"{{{W}}}rStyle")
    ref_style.set(f"{{{W}}}val", "CommentReference")
    ref_ref = lxml_etree.SubElement(ref_run, f"{{{W}}}commentReference")
    ref_ref.set(f"{{{W}}}id", str(comment_id))

    # Now rearrange: remove the three new elements, reinsert at correct positions
    elem.remove(range_start)
    elem.remove(range_end)
    elem.remove(ref_run)

    # Find the index of the first run
    children = list(elem)
    first_run_idx = children.index(runs[0])

    # Insert rangeStart before first run
    elem.insert(first_run_idx, range_start)

    # Refresh children list after insert
    children = list(elem)
    last_run_idx = children.index(runs[-1])

    # Insert rangeEnd after last run
    elem.insert(last_run_idx + 1, range_end)

    # Insert reference run after rangeEnd
    children = list(elem)
    range_end_idx = children.index(range_end)
    elem.insert(range_end_idx + 1, ref_run)


def _patch_zip_with_comments(temp_docx, output_path, comment_data):
    """Patch the saved DOCX ZIP to add comment parts."""
    # Build comments.xml
    comments_xml = _build_comments_xml(comment_data)

    # Build commentsExtended.xml
    comments_ext_xml = _build_comments_extended_xml(comment_data)

    # Read existing ZIP, write new ZIP with additions
    with zipfile.ZipFile(temp_docx, "r") as zin:
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                data = zin.read(item)

                if item == "[Content_Types].xml":
                    data = _patch_content_types(data)
                elif item == "word/_rels/document.xml.rels":
                    data = _patch_relationships(data)

                zout.writestr(item, data)

            # Add new parts
            zout.writestr("word/comments.xml", comments_xml)
            zout.writestr("word/commentsExtended.xml", comments_ext_xml)


def _build_comments_xml(comment_data):
    """Build word/comments.xml content."""
    lines = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        f'<w:comments xmlns:w="{W}" xmlns:r="{R}" xmlns:w14="{W14}">',
    ]

    for c in comment_data:
        lines.append(
            f'  <w:comment w:id="{c["id"]}" w:author="{c["author"]}" '
            f'w:date="{c["date"]}" w:initials="{c["initials"]}">'
        )
        lines.append(
            f'    <w:p w14:paraId="{c["para_id"]}">'
            f'<w:r><w:t>{_xml_escape(c["text"])}</w:t></w:r>'
            f'</w:p>'
        )
        lines.append('  </w:comment>')

    lines.append('</w:comments>')
    return "\n".join(lines)


def _build_comments_extended_xml(comment_data):
    """Build word/commentsExtended.xml content."""
    lines = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        f'<w15:commentsEx xmlns:w15="{W15}" xmlns:mc="{MC}" mc:Ignorable="w15">',
    ]

    for c in comment_data:
        done = "1" if c["resolved"] else "0"
        if c.get("reply_to") is not None:
            # Find the parent comment's paraId
            parent_idx = c["reply_to"]
            if 0 <= parent_idx < len(comment_data):
                parent_para_id = comment_data[parent_idx]["para_id"]
                lines.append(
                    f'  <w15:commentEx w15:paraId="{c["para_id"]}" '
                    f'w15:paraIdParent="{parent_para_id}" w15:done="{done}"/>'
                )
            else:
                lines.append(f'  <w15:commentEx w15:paraId="{c["para_id"]}" w15:done="{done}"/>')
        else:
            lines.append(f'  <w15:commentEx w15:paraId="{c["para_id"]}" w15:done="{done}"/>')

    lines.append('</w15:commentsEx>')
    return "\n".join(lines)


def _patch_content_types(data):
    """Add comment content types to [Content_Types].xml."""
    content = data.decode("utf-8")

    comments_ct = 'PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"'
    comments_ext_ct = 'PartName="/word/commentsExtended.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml"'

    if "comments.xml" not in content:
        content = content.replace("</Types>", f'  <Override {comments_ct}/>\n</Types>')
    if "commentsExtended.xml" not in content:
        content = content.replace("</Types>", f'  <Override {comments_ext_ct}/>\n</Types>')

    return content.encode("utf-8")


def _patch_relationships(data):
    """Add comment relationships to word/_rels/document.xml.rels."""
    content = data.decode("utf-8")

    # Find next available rId
    import re
    rids = [int(m) for m in re.findall(r'Id="rId(\d+)"', content)]
    next_rid = max(rids) + 1 if rids else 1

    comments_rel = (
        f'<Relationship Id="rId{next_rid}" '
        f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" '
        f'Target="comments.xml"/>'
    )
    comments_ext_rel = (
        f'<Relationship Id="rId{next_rid + 1}" '
        f'Type="http://schemas.microsoft.com/office/2011/relationships/commentsExtended" '
        f'Target="commentsExtended.xml"/>'
    )

    if "relationships/comments" not in content:
        content = content.replace("</Relationships>", f"  {comments_rel}\n</Relationships>")
    if "relationships/commentsExtended" not in content:
        content = content.replace("</Relationships>", f"  {comments_ext_rel}\n</Relationships>")

    return content.encode("utf-8")


def _xml_escape(text):
    """Escape text for XML content."""
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&apos;"))


def main():
    parser = argparse.ArgumentParser(
        description="Add comments to DOCX from JSON manifest",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("input", type=Path, help="Input DOCX file")
    parser.add_argument("output", type=Path, help="Output DOCX file")
    parser.add_argument("--comments", required=True, type=Path,
                        help="JSON file with comment definitions")
    parser.add_argument("--author", default="RIO AI", help="Comment author (default: RIO AI)")
    parser.add_argument("--initials", default="RI", help="Author initials (default: RI)")

    args = parser.parse_args()

    if not args.input.is_file():
        print(f"Error: File not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    if not args.comments.is_file():
        print(f"Error: Comments file not found: {args.comments}", file=sys.stderr)
        sys.exit(1)

    add_comments(
        args.input, args.output, args.comments,
        author=args.author, initials=args.initials,
    )


if __name__ == "__main__":
    main()
