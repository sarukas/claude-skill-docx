#!/usr/bin/env python3
"""
Self-contained Markdown to DOCX converter.

Converts Markdown files to professional Word documents without Pandoc.
Uses mistune for Markdown parsing and python-docx for DOCX generation.

Usage:
    python md_to_docx_py.py input.md output.docx [--template template.docx]

Dependencies (pip install):
    mistune, python-docx, Pillow, requests
"""

from __future__ import annotations

import argparse
import base64
import os
import re
import subprocess
import sys
import tempfile
from io import BytesIO
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from typing import Any

try:
    import mistune
except ImportError:
    print("Installing mistune...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "mistune"])
    import mistune

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import requests
except ImportError:
    requests = None


# ---------------------------------------------------------------------------
# Style defaults & configuration
# ---------------------------------------------------------------------------
_STYLE_DEFAULTS = {
    "font_body": "Arial",
    "font_heading": "Arial",
    "font_code": "Consolas",
    "font_size": "10.5",
    "color_heading": "2D3B4D",
    "color_body": "333333",
    "table_header_bg": "D5E8F0",
    "table_header_text": "2D3B4D",
    "table_alt_row": "F2F2F2",
    "table_border": "CCCCCC",
    "table_border_size": "4",
    "table_cell_margin": "28",
    "table_font_size": "9.5",
    "table_banded_rows": "true",
    "code_bg": "F5F5F5",
    "code_font_size": "9",
}

# Module-level style variables (set by _apply_style, used throughout)
FONT_BODY = "Arial"
FONT_HEADING = "Arial"
FONT_CODE = "Consolas"
FONT_SIZE_BODY = 10.5
FONT_SIZE_TABLE = 9.5
FONT_SIZE_CODE = 9.0
COLOR_HEADING = RGBColor(0x2D, 0x3B, 0x4D)
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)
COLOR_LINK = RGBColor(0x05, 0x63, 0xC1)
COLOR_TABLE_HEADER_BG = "D5E8F0"
COLOR_TABLE_HEADER_TEXT = RGBColor(0x2D, 0x3B, 0x4D)
COLOR_TABLE_ALT_ROW = "F2F2F2"
COLOR_TABLE_BORDER = "CCCCCC"
COLOR_CODE_BG = "F5F5F5"
TABLE_BORDER_SIZE = 4
TABLE_CELL_MARGIN = 28
TABLE_BANDED_ROWS = True
HEADING_SIZES = {1: 20, 2: 16, 3: 14, 4: 12, 5: 11, 6: 10.5}


def _hex_to_rgb(hex_str: str) -> RGBColor:
    """Convert hex color string (e.g. '2D3B4D') to RGBColor."""
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


_DOCX_STYLE_RE = re.compile(r"<!--\s*docx-style\s*\n(.*?)-->", re.DOTALL)


def parse_docx_style(text: str) -> dict:
    """Extract docx-style configuration from HTML comment in markdown."""
    match = _DOCX_STYLE_RE.search(text)
    if not match:
        return {}
    config = {}
    for line in match.group(1).strip().split("\n"):
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if ":" in line:
            key, value = line.split(":", 1)
            key = key.strip().lower()
            value = value.strip().strip('"').strip("'")
            if key in _STYLE_DEFAULTS:
                config[key] = value
    return config


def strip_docx_style_comment(text: str) -> str:
    """Remove the docx-style comment from markdown content."""
    return _DOCX_STYLE_RE.sub("", text)


def parse_style_file(path: str) -> dict:
    """Parse a style configuration file (same key: value format)."""
    config = {}
    for line in Path(path).read_text(encoding="utf-8").strip().split("\n"):
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if ":" in line:
            key, value = line.split(":", 1)
            key = key.strip().lower()
            value = value.strip().strip('"').strip("'")
            if key in _STYLE_DEFAULTS:
                config[key] = value
    return config


def _open_template(template_path: str) -> "Document":
    """Open a .docx or .dotx template as a python-docx Document.

    Word .dotx templates are structurally identical to .docx files except for
    a single content-type string in [Content_Types].xml.  python-docx rejects
    .dotx files, so we monkey-patch the content type and convert in-memory.
    """
    if template_path.lower().endswith(('.dotx', '.dotm')):
        from docx.opc.constants import CONTENT_TYPE as CT
        from docx.opc.package import OpcPackage
        from docx.opc.part import PartFactory
        from docx.parts.document import DocumentPart

        # Register template content types so PartFactory creates DocumentPart
        for ct in (
            'application/vnd.openxmlformats-officedocument'
            '.wordprocessingml.template.main+xml',
            'application/vnd.ms-word.template.macroEnabledTemplate.main+xml',
        ):
            PartFactory.part_type_for[ct] = DocumentPart

        package = OpcPackage.open(template_path)
        package.main_document_part._content_type = CT.WML_DOCUMENT_MAIN  # type: ignore[attr-defined]

        buf = BytesIO()
        package.save(buf)
        buf.seek(0)
        return Document(buf)
    return Document(template_path)


def extract_style_from_template(template_path: str) -> dict:
    """Extract style configuration from a DOCX template's actual formatting.

    Reads docDefaults, heading styles, table styles, and theme colors to
    produce a dict compatible with _STYLE_DEFAULTS keys.  The result can be
    used as a base style layer that is then overridden by --style / inline /
    CLI flags.
    """
    from lxml import etree  # noqa: F811 – local import to keep top-level light

    doc = _open_template(template_path)
    config: dict[str, str] = {}
    wns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    # --- 1. Document defaults (fonts, body size) ----------------------------
    styles_el = doc.styles.element
    doc_defaults = styles_el.find(f".//{wns}docDefaults")
    if doc_defaults is not None:
        rpr = doc_defaults.find(f"{wns}rPrDefault/{wns}rPr")
        if rpr is not None:
            rfonts = rpr.find(f"{wns}rFonts")
            if rfonts is not None:
                ascii_font = rfonts.get(f"{wns}ascii")
                if ascii_font:
                    config["font_body"] = ascii_font
                    config["font_heading"] = ascii_font  # may be overridden below
            sz = rpr.find(f"{wns}sz")
            if sz is not None:
                half_pts = int(sz.get(f"{wns}val"))
                config["font_size"] = str(half_pts / 2)
            # Body text color from defaults
            color_el = rpr.find(f"{wns}color")
            if color_el is not None:
                val = color_el.get(f"{wns}val")
                if val and val.lower() != "auto":
                    config["color_body"] = val

    # --- 2. Normal style (body color fallback) ------------------------------
    try:
        normal = doc.styles["Normal"]
        n_el = normal.element.find(f"{wns}rPr")
        if n_el is not None:
            color_el = n_el.find(f"{wns}color")
            if color_el is not None:
                val = color_el.get(f"{wns}val")
                if val and val.lower() != "auto":
                    config["color_body"] = val
    except KeyError:
        pass

    # --- 3. Theme colors (fallback for headings / body) ---------------------
    theme_dk1 = None
    try:
        theme_part = doc.part.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme"
        )
        theme_xml = etree.fromstring(theme_part.blob)
        ans = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
        # dk1 = default dark text color
        for dk1_el in theme_xml.findall(".//a:clrScheme/a:dk1/*", ans):
            theme_dk1 = dk1_el.get("lastClr") or dk1_el.get("val")
            break
    except Exception:
        pass

    # Apply theme fallbacks where no explicit color was found
    if "color_body" not in config:
        config["color_body"] = theme_dk1 or "000000"
    if "color_heading" not in config:
        config["color_heading"] = theme_dk1 or "000000"

    # --- 4. Heading styles (font, color override) ---------------------------
    for level in [1, 2, 3]:
        try:
            h = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        h_rpr = h.element.find(f"{wns}rPr")
        if h_rpr is None:
            continue
        # Explicit heading font
        rfonts = h_rpr.find(f"{wns}rFonts")
        if rfonts is not None:
            ascii_font = rfonts.get(f"{wns}ascii")
            if ascii_font:
                config["font_heading"] = ascii_font
        # Explicit heading color
        color_el = h_rpr.find(f"{wns}color")
        if color_el is not None:
            val = color_el.get(f"{wns}val")
            if val and val.lower() != "auto":
                config["color_heading"] = val
        break  # first available heading is enough

    # --- 5. Table styles (header bg, header text, banded rows) --------------
    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.TABLE:
            continue
        el = style.element
        tbl_style_prs = el.findall(f"{wns}tblStylePr")
        if not tbl_style_prs:
            continue
        for tsp in tbl_style_prs:
            tsp_type = tsp.get(f"{wns}type")
            if tsp_type != "firstRow":
                continue
            # Header background
            shd = tsp.find(f".//{wns}shd")
            if shd is not None:
                fill = shd.get(f"{wns}fill")
                if fill and fill.lower() not in ("auto", "ffffff"):
                    config["table_header_bg"] = fill
            # Header text color
            rpr = tsp.find(f".//{wns}rPr/{wns}color")
            if rpr is not None:
                val = rpr.get(f"{wns}val")
                if val:
                    config["table_header_text"] = val
            if "table_header_bg" in config:
                break
        if "table_header_bg" in config:
            break

    # --- 6. Table paragraph style (font size) -------------------------------
    for sname in ("Table1", "Table Contents", "Table Text"):
        try:
            ts = doc.styles[sname]
        except KeyError:
            continue
        ts_rpr = ts.element.find(f".//{wns}rPr/{wns}sz")
        if ts_rpr is not None:
            half_pts = int(ts_rpr.get(f"{wns}val"))
            config["table_font_size"] = str(half_pts / 2)
            break

    return config


def _apply_style(overrides: dict):
    """Apply style configuration. Starts from defaults, overlays overrides."""
    global FONT_BODY, FONT_HEADING, FONT_CODE, FONT_SIZE_BODY, FONT_SIZE_TABLE, FONT_SIZE_CODE
    global COLOR_HEADING, COLOR_BODY, COLOR_LINK
    global COLOR_TABLE_HEADER_BG, COLOR_TABLE_HEADER_TEXT, COLOR_TABLE_ALT_ROW
    global COLOR_TABLE_BORDER, COLOR_CODE_BG
    global TABLE_BORDER_SIZE, TABLE_CELL_MARGIN, TABLE_BANDED_ROWS

    cfg = dict(_STYLE_DEFAULTS)
    cfg.update(overrides)

    FONT_BODY = cfg["font_body"]
    FONT_HEADING = cfg["font_heading"]
    FONT_CODE = cfg["font_code"]
    FONT_SIZE_BODY = float(cfg["font_size"])
    FONT_SIZE_TABLE = float(cfg["table_font_size"])
    FONT_SIZE_CODE = float(cfg["code_font_size"])
    COLOR_HEADING = _hex_to_rgb(cfg["color_heading"])
    COLOR_BODY = _hex_to_rgb(cfg["color_body"])
    COLOR_TABLE_HEADER_BG = cfg["table_header_bg"]
    COLOR_TABLE_HEADER_TEXT = _hex_to_rgb(cfg["table_header_text"])
    COLOR_TABLE_ALT_ROW = cfg["table_alt_row"]
    COLOR_TABLE_BORDER = cfg["table_border"]
    COLOR_CODE_BG = cfg["code_bg"]
    TABLE_BORDER_SIZE = int(cfg["table_border_size"])
    TABLE_CELL_MARGIN = int(cfg["table_cell_margin"])
    TABLE_BANDED_ROWS = cfg["table_banded_rows"].lower() in ("true", "yes", "1")


# ---------------------------------------------------------------------------
# Mermaid helpers  (reused from existing converter)
# ---------------------------------------------------------------------------
def generate_mermaid_url(code: str) -> str | None:
    try:
        encoded = base64.urlsafe_b64encode(code.encode()).decode().rstrip("=")
        url = f"https://mermaid.ink/img/{encoded}"
        return url if len(url) <= 2000 else None
    except Exception:
        return None


def check_mermaid_cli():
    paths = [
        "mmdc",
        os.path.expandvars(r"%APPDATA%\npm\mmdc.cmd"),
        os.path.expandvars(r"%APPDATA%\npm\mmdc"),
    ]
    for p in paths:
        try:
            r = subprocess.run([p, "--version"], capture_output=True, text=True, timeout=10, shell=True)
            if r.returncode == 0:
                return p
        except Exception:
            continue
    return None


def render_mermaid(code: str, out_path: str) -> bool:
    # 1. Try URL
    url = generate_mermaid_url(code)
    if url and requests:
        try:
            resp = requests.get(url, timeout=30)
            resp.raise_for_status()
            Path(out_path).write_bytes(resp.content)
            print(f"  Mermaid: downloaded via URL ({len(resp.content)} bytes)")
            return True
        except Exception as e:
            print(f"  Mermaid URL failed: {e}")

    # 2. Try local CLI
    cli = check_mermaid_cli()
    if cli:
        mmd = out_path.replace(".png", ".mmd")
        Path(mmd).write_text(code, encoding="utf-8")
        try:
            r = subprocess.run(
                [cli, "-i", mmd, "-o", out_path, "-b", "transparent"],
                capture_output=True, text=True, timeout=60, shell=True,
            )
            try:
                os.remove(mmd)
            except OSError:
                pass
            if r.returncode == 0 and os.path.exists(out_path):
                print(f"  Mermaid: rendered locally ({os.path.getsize(out_path)} bytes)")
                return True
            print(f"  Mermaid CLI error: {r.stderr or r.stdout}")
        except Exception as e:
            print(f"  Mermaid CLI exception: {e}")

    return False


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------
class DocxBuilder:
    """Accumulates python-docx elements from parsed Markdown tokens."""

    def __init__(self, doc: Document, input_dir: str, temp_dir: str):
        self.doc = doc
        self.input_dir = input_dir
        self.temp_dir = temp_dir
        self.mermaid_counter = 0
        self._setup_styles()

    # -- style helpers -------------------------------------------------------
    def _setup_styles(self):
        """Configure document-level styles (only when no template is used)."""
        style = self.doc.styles
        # Normal
        normal = style["Normal"]
        normal.font.name = FONT_BODY
        normal.font.size = Pt(FONT_SIZE_BODY)
        normal.font.color.rgb = COLOR_BODY
        pf = normal.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15

        # Headings
        for level in range(1, 7):
            sname = f"Heading {level}"
            try:
                h = style[sname]
            except KeyError:
                h = style.add_style(sname, WD_STYLE_TYPE.PARAGRAPH)
            h.font.name = FONT_HEADING
            h.font.size = Pt(HEADING_SIZES[level])
            h.font.bold = True
            h.font.color.rgb = COLOR_HEADING
            h.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True

        # Code style (character)
        if "Code Char" not in [s.name for s in style]:
            cs = style.add_style("Code Char", WD_STYLE_TYPE.CHARACTER)
            cs.font.name = FONT_CODE
            cs.font.size = Pt(FONT_SIZE_CODE)
            cs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # -- inline helpers ------------------------------------------------------
    @staticmethod
    def _set_shading(element, color: str):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
        element.append(shading)

    def _add_inline(self, paragraph, tokens):
        """Render inline tokens (text, bold, italic, code, link, image, etc.) into a paragraph."""
        if isinstance(tokens, str):
            paragraph.add_run(tokens)
            return
        if not isinstance(tokens, list):
            tokens = [tokens]
        for tok in tokens:
            if isinstance(tok, str):
                paragraph.add_run(tok)
                continue
            tp = tok.get("type", "")
            children = tok.get("children", [])
            raw = tok.get("raw", tok.get("text", ""))

            if tp == "text":
                paragraph.add_run(tok.get("raw", tok.get("text", tok.get("children", ""))))
            elif tp == "softbreak":
                paragraph.add_run("\n")
            elif tp == "linebreak":
                run = paragraph.add_run()
                run.add_break()
            elif tp == "strong":
                run = paragraph.add_run(self._flatten_text(children))
                run.bold = True
            elif tp == "emphasis":
                run = paragraph.add_run(self._flatten_text(children))
                run.italic = True
            elif tp == "strikethrough":
                run = paragraph.add_run(self._flatten_text(children))
                run.font.strike = True
            elif tp == "codespan":
                run = paragraph.add_run(raw)
                run.font.name = FONT_CODE
                run.font.size = Pt(FONT_SIZE_CODE)
                self._set_shading(run._element.get_or_add_rPr(), COLOR_CODE_BG)
            elif tp == "link":
                link_url = (tok.get("attrs") or {}).get("url", "") or tok.get("link", "")
                self._add_hyperlink(paragraph, link_url, self._flatten_text(children))
            elif tp == "image":
                self._add_image_inline(paragraph, tok.get("src", ""), tok.get("alt", ""))
            else:
                # fallback – just dump text
                paragraph.add_run(self._flatten_text(children) if children else raw)

    def _flatten_text(self, tokens) -> str:
        """Recursively extract plain text from token tree."""
        if isinstance(tokens, str):
            return tokens
        if isinstance(tokens, dict):
            children = tokens.get("children", tokens.get("raw", tokens.get("text", "")))
            return self._flatten_text(children)
        if isinstance(tokens, list):
            return "".join(self._flatten_text(t) for t in tokens)
        return str(tokens)

    def _add_hyperlink(self, paragraph, url: str, text: str):
        """Insert a clickable hyperlink into a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = parse_xml(
            f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f'<w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr>'
            f"<w:t>{self._escape_xml(text)}</w:t></w:r></w:hyperlink>"
        )
        paragraph._element.append(hyperlink)

    @staticmethod
    def _escape_xml(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

    # -- image helpers -------------------------------------------------------
    def _add_image_inline(self, paragraph, src: str, alt: str = ""):
        path = self._resolve_image(src)
        if path and os.path.isfile(path):
            try:
                width, height = self._image_fit(path)
                paragraph.add_run().add_picture(path, width=width, height=height)
            except Exception as e:
                paragraph.add_run(f"[Image: {alt or src} - {e}]").italic = True
        else:
            paragraph.add_run(f"[Image: {alt or src}]").italic = True

    def _add_image_block(self, src: str, alt: str = ""):
        path = self._resolve_image(src)
        if path and os.path.isfile(path):
            try:
                width, height = self._image_fit(path)
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run().add_picture(path, width=width, height=height)
            except Exception as e:
                p = self.doc.add_paragraph()
                run = p.add_run(f"[Image: {alt or src} - {e}]")
                run.italic = True
        else:
            p = self.doc.add_paragraph()
            run = p.add_run(f"[Image: {alt or src}]")
            run.italic = True

    def _resolve_image(self, src: str) -> str | None:
        if src.startswith(("http://", "https://")) and requests:
            try:
                resp = requests.get(src, timeout=15)
                resp.raise_for_status()
                ext = Path(src).suffix or ".png"
                tmp = os.path.join(self.temp_dir, f"img_{hash(src)}{ext}")
                Path(tmp).write_bytes(resp.content)
                return tmp
            except Exception:
                return None
        # local path
        p = Path(src)
        if not p.is_absolute():
            p = Path(self.input_dir) / p
        return str(p) if p.is_file() else None

    @staticmethod
    def _image_fit(path: str):
        """Return (width, height) in Inches, constrained to A4 content area minus 15%."""
        # A4 with 1" margins: 6.27" x 9.69"
        # Width at 100% of content area; height at 85% for header/footer room
        max_w = 6.27
        max_h = 9.69 * 0.85  # ~8.24"
        if Image:
            try:
                with Image.open(path) as img:
                    w_px, h_px = img.size
                    dpi = img.info.get("dpi", (96, 96))[0] or 96
                    w_in = w_px / dpi
                    h_in = h_px / dpi
                    scale = min(1.0, max_w / w_in, max_h / h_in)
                    return Inches(w_in * scale), Inches(h_in * scale)
            except Exception:
                pass
        return Inches(5), Inches(3.5)

    # -- table helpers -------------------------------------------------------
    def _add_table(self, header_tokens: list, body_tokens: list[list]):
        """Build a professional table. header_tokens and body cell items are
        raw mistune inline token lists so formatting (bold, code, etc.) is preserved."""
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

        ncols = len(header_tokens)
        nrows = 1 + len(body_tokens)
        table = self.doc.add_table(rows=nrows, cols=ncols)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass  # Template may not include "Table Grid" – use default style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Column widths – distribute evenly within ~6 inches
        col_width = Inches(6.0 / ncols) if ncols else Inches(3)
        for col in table.columns:
            col.width = col_width

        # Cell margin: 0.05 cm ≈ 28 twips on all sides
        self._set_table_cell_margins(table, TABLE_CELL_MARGIN)

        def _format_cell(cell, tokens, is_header=False, shade_color=None):
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_inline(p, tokens)
            for run in p.runs:
                run.font.size = Pt(FONT_SIZE_TABLE)
                run.font.name = FONT_BODY
                if is_header:
                    run.bold = True
                    run.font.color.rgb = COLOR_TABLE_HEADER_TEXT
            if shade_color:
                self._shade_cell(cell, shade_color)

        # Header row
        hdr = table.rows[0]
        for i, cell in enumerate(hdr.cells):
            _format_cell(cell, header_tokens[i], is_header=True, shade_color=COLOR_TABLE_HEADER_BG)

        # Body rows – banded rows only (alternating), no banded columns
        for r_idx, row_data in enumerate(body_tokens):
            row = table.rows[r_idx + 1]
            for c_idx, cell in enumerate(row.cells):
                tokens = row_data[c_idx] if c_idx < len(row_data) else []
                shade = COLOR_TABLE_ALT_ROW if TABLE_BANDED_ROWS and r_idx % 2 == 1 else None
                _format_cell(cell, tokens, shade_color=shade)

        # Borders
        self._set_table_borders(table)
        # Spacing after table
        self.doc.add_paragraph()

    @staticmethod
    def _table_align(alignments, idx):
        if not alignments or idx >= len(alignments):
            return WD_ALIGN_PARAGRAPH.LEFT
        a = (alignments[idx] or "").lower()
        if "center" in a:
            return WD_ALIGN_PARAGRAPH.CENTER
        if "right" in a:
            return WD_ALIGN_PARAGRAPH.RIGHT
        return WD_ALIGN_PARAGRAPH.LEFT

    @staticmethod
    def _shade_cell(cell, color: str):
        tc_pr = cell._element.get_or_add_tcPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
        tc_pr.append(shading)

    @staticmethod
    def _set_table_borders(table):
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f"<w:tblPr {nsdecls('w')}/>")
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="{TABLE_BORDER_SIZE}" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
            f'<w:left w:val="single" w:sz="{TABLE_BORDER_SIZE}" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
            f'<w:bottom w:val="single" w:sz="{TABLE_BORDER_SIZE}" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
            f'<w:right w:val="single" w:sz="{TABLE_BORDER_SIZE}" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
            f'<w:insideH w:val="single" w:sz="{TABLE_BORDER_SIZE}" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
            f'<w:insideV w:val="single" w:sz="{TABLE_BORDER_SIZE}" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
            f"</w:tblBorders>"
        )
        tbl_pr.append(borders)

    @staticmethod
    def _set_table_cell_margins(table, margin_twips: int = 28):
        """Set uniform cell margins (0.05 cm ≈ 28 twips) on the table."""
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = parse_xml(f"<w:tblPr {nsdecls('w')}/>")
            table._tbl.append(tbl_pr)
        m = str(margin_twips)
        cell_mar = parse_xml(
            f'<w:tblCellMar {nsdecls("w")}>'
            f'<w:top w:w="{m}" w:type="dxa"/>'
            f'<w:left w:w="{m}" w:type="dxa"/>'
            f'<w:bottom w:w="{m}" w:type="dxa"/>'
            f'<w:right w:w="{m}" w:type="dxa"/>'
            f"</w:tblCellMar>"
        )
        tbl_pr.append(cell_mar)

    # -- code block ----------------------------------------------------------
    def _add_code_block(self, code: str, _lang: str = ""):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(4)
        pf.space_after = Pt(4)
        pf.left_indent = Inches(0.25)
        # background shading on paragraph
        pPr = p._element.get_or_add_pPr()
        self._set_shading(pPr, COLOR_CODE_BG)
        for i, line in enumerate(code.rstrip("\n").split("\n")):
            if i > 0:
                run = p.add_run()
                run.add_break()
            run = p.add_run(line)
            run.font.name = FONT_CODE
            run.font.size = Pt(FONT_SIZE_CODE)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # -- blockquote ----------------------------------------------------------
    def _add_blockquote(self, children):
        """Render blockquote with left indent."""
        for tok in children:
            tp = tok.get("type", "")
            if tp == "paragraph":
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                # Add left border via XML
                pPr = p._element.get_or_add_pPr()
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'<w:left w:val="single" w:sz="12" w:space="8" w:color="AAAAAA"/>'
                    f"</w:pBdr>"
                )
                pPr.append(pBdr)
                self._add_inline(p, tok.get("children", []))
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            elif tp == "block_code":
                self._add_code_block(tok.get("raw", tok.get("text", "")), tok.get("attrs", {}).get("info", ""))
            else:
                self._render_token(tok)

    # -- lists ---------------------------------------------------------------
    def _add_list(self, token, level: int = 0):
        ordered = token.get("attrs", {}).get("ordered", False)
        children = token.get("children", [])
        counter = token.get("attrs", {}).get("start", 1) or 1

        for item in children:
            if item.get("type") != "list_item":
                continue
            item_children = item.get("children", [])
            first_para = True
            for child in item_children:
                tp = child.get("type", "")
                # mistune v3 uses "block_text" for tight lists, "paragraph" for loose
                if tp in ("paragraph", "block_text"):
                    p = self.doc.add_paragraph()
                    indent = Inches(0.25 + level * 0.25)
                    hanging = Inches(0.25)
                    p.paragraph_format.left_indent = indent + hanging
                    p.paragraph_format.first_line_indent = -hanging
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)

                    if first_para:
                        if ordered:
                            prefix = f"{counter}. "
                            counter += 1
                        else:
                            bullets = ["\u2022", "\u25E6", "\u25AA"]
                            prefix = bullets[min(level, len(bullets) - 1)] + " "
                        run = p.add_run(prefix)
                        run.font.name = FONT_BODY
                        run.font.size = Pt(10.5)
                        first_para = False

                    self._add_inline(p, child.get("children", []))
                elif tp == "list":
                    self._add_list(child, level + 1)
                elif tp == "block_code":
                    self._add_code_block(child.get("raw", child.get("text", "")))

    # -- main render dispatch ------------------------------------------------
    def render_tokens(self, tokens: list):
        for tok in tokens:
            self._render_token(tok)

    def _render_token(self, tok: dict):
        tp = tok.get("type", "")

        if tp == "heading":
            level = tok.get("attrs", {}).get("level", 1)
            text = self._flatten_text(tok.get("children", []))
            p = self.doc.add_heading(text, level=min(level, 6))
            # Ensure heading styling is applied
            for run in p.runs:
                run.font.name = FONT_HEADING
                run.font.color.rgb = COLOR_HEADING

        elif tp == "paragraph":
            children = tok.get("children", [])
            # Check if sole child is an image
            if len(children) == 1 and isinstance(children[0], dict) and children[0].get("type") == "image":
                img = children[0]
                self._add_image_block(img.get("src", ""), img.get("alt", ""))
            else:
                p = self.doc.add_paragraph()
                self._add_inline(p, children)

        elif tp == "block_code":
            info = tok.get("attrs", {}).get("info", "") or ""
            raw = tok.get("raw", tok.get("text", ""))
            if info.strip().lower() == "mermaid":
                self._handle_mermaid(raw)
            else:
                self._add_code_block(raw, info)

        elif tp == "table":
            self._handle_table(tok)

        elif tp == "list":
            self._add_list(tok)

        elif tp == "block_quote":
            self._add_blockquote(tok.get("children", []))

        elif tp == "thematic_break":
            # Horizontal rule -> page break
            self.doc.add_page_break()

        elif tp == "blank_line":
            pass  # ignore

        elif tp == "block_html":
            raw = tok.get("raw", "")
            if raw.strip():
                p = self.doc.add_paragraph()
                run = p.add_run(raw.strip())
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        else:
            # Unknown block – try to render children
            children = tok.get("children")
            if children and isinstance(children, list):
                self.render_tokens(children)

    # -- table token handling ------------------------------------------------
    def _handle_table(self, tok: dict):
        """Handle mistune table token -> professional DOCX table.

        Mistune v3 AST for tables:
          table -> [table_head, table_body]
          table_head -> [table_cell, table_cell, ...]   (cells directly, no row wrapper)
          table_body -> [table_row, table_row, ...]
          table_row  -> [table_cell, table_cell, ...]

        Passes raw inline token lists (not flattened text) so bold/italic/code
        formatting inside cells is preserved.
        """
        children = tok.get("children", [])
        header_tokens = []   # list of token-lists, one per header cell
        body_tokens = []     # list of rows, each row is list of token-lists

        for child in children:
            ctype = child.get("type", "")
            if ctype == "table_head":
                for item in child.get("children", []):
                    itype = item.get("type", "")
                    if itype == "table_cell":
                        header_tokens.append(item.get("children", []))
                    elif itype == "table_row":
                        for cell in item.get("children", []):
                            header_tokens.append(cell.get("children", []))
            elif ctype == "table_body":
                for row in child.get("children", []):
                    if row.get("type") == "table_row":
                        row_data = []
                        for cell in row.get("children", []):
                            row_data.append(cell.get("children", []))
                        body_tokens.append(row_data)

        if header_tokens:
            self._add_table(header_tokens, body_tokens)

    # -- mermaid -------------------------------------------------------------
    def _handle_mermaid(self, code: str):
        self.mermaid_counter += 1
        out_path = os.path.join(self.temp_dir, f"mermaid_{self.mermaid_counter}.png")
        if render_mermaid(code.strip(), out_path):
            self._add_image_block(out_path, f"Mermaid Diagram {self.mermaid_counter}")
            self.doc.add_paragraph()  # spacer for natural page-break flow
        else:
            p = self.doc.add_paragraph()
            run = p.add_run(f"[Mermaid diagram {self.mermaid_counter} could not be rendered]")
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ---------------------------------------------------------------------------
# YAML front-matter stripper
# ---------------------------------------------------------------------------
_FRONT_MATTER_RE = re.compile(r"^---\s*\n.*?\n---\s*\n", re.DOTALL)


def strip_front_matter(text: str) -> str:
    return _FRONT_MATTER_RE.sub("", text, count=1)


# ---------------------------------------------------------------------------
# Title page, TOC, footer helpers
# ---------------------------------------------------------------------------
def _add_title_page(doc, builder, title: str, preamble_tokens: list,
                    date: str | None = None):
    """Add a centred title page with optional preamble content, then page break."""
    # Push title towards vertical centre: add spacer paragraphs
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = FONT_HEADING
    run.font.color.rgb = COLOR_HEADING

    # Date (optional, from --date)
    if date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(date)
        run.font.size = Pt(14)
        run.font.name = FONT_BODY
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Render preamble tokens (paragraphs between H1 and first ---) centred
    if preamble_tokens:
        doc.add_paragraph()  # spacer
        for tok in preamble_tokens:
            if tok.get("type") == "paragraph":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                builder._add_inline(p, tok.get("children", []))
            elif tok.get("type") == "blank_line":
                continue  # skip blank lines in preamble
            else:
                builder._render_token(tok)

    # Page break (new section for body content)
    doc.add_page_break()


def _add_toc(doc):
    """Insert a Table of Contents field that Word will populate on open."""
    # Set updateFields so Word prompts to refresh all fields (incl. TOC) on open
    settings = doc.settings.element
    update_fields = parse_xml(f'<w:updateFields {nsdecls("w")} w:val="true"/>')
    settings.append(update_fields)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = FONT_HEADING
    run.font.color.rgb = COLOR_HEADING

    # Insert TOC field via XML
    p = doc.add_paragraph()
    fld_begin = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:fldChar w:fldCharType="begin"/>'
        f'</w:r>'
    )
    fld_code = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
        f'</w:r>'
    )
    fld_separate = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:fldChar w:fldCharType="separate"/>'
        f'</w:r>'
    )
    fld_placeholder = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:t>Right-click and select "Update Field" to generate table of contents.</w:t>'
        f'</w:r>'
    )
    fld_end = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:fldChar w:fldCharType="end"/>'
        f'</w:r>'
    )
    p._element.append(fld_begin)
    p._element.append(fld_code)
    p._element.append(fld_separate)
    p._element.append(fld_placeholder)
    p._element.append(fld_end)

    doc.add_page_break()


def _setup_footer(doc, pagination: bool = True, copyright_text: str | None = None):
    """Configure footer on all sections with optional page numbers and copyright."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        # Clear existing footer
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)

        parts = []
        if copyright_text:
            parts.append(copyright_text)
        if pagination:
            parts.append(None)  # placeholder for page number

        if not parts:
            continue

        p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/>'
                      f'<w:rPr><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
                      f'</w:pPr></w:p>')

        _rpr = (f'<w:rPr><w:rFonts w:ascii="{FONT_BODY}" w:hAnsi="{FONT_BODY}"/>'
                f'<w:sz w:val="16"/><w:szCs w:val="16"/><w:color w:val="AAAAAA"/></w:rPr>')

        if copyright_text and pagination:
            # "Copyright  |  Page X"
            for text_piece in [copyright_text, "  |  Page "]:
                r = parse_xml(
                    f'<w:r {nsdecls("w")}>{_rpr}'
                    f'<w:t xml:space="preserve">{text_piece}</w:t></w:r>'
                )
                p.append(r)
            # Page number field – all three runs need matching rPr
            r = parse_xml(
                f'<w:r {nsdecls("w")}>{_rpr}'
                f'<w:fldChar w:fldCharType="begin"/></w:r>'
            )
            p.append(r)
            r = parse_xml(
                f'<w:r {nsdecls("w")}>{_rpr}'
                f'<w:instrText> PAGE </w:instrText></w:r>'
            )
            p.append(r)
            r = parse_xml(
                f'<w:r {nsdecls("w")}>{_rpr}'
                f'<w:fldChar w:fldCharType="end"/></w:r>'
            )
            p.append(r)
        elif pagination:
            # "Page X" only
            r = parse_xml(
                f'<w:r {nsdecls("w")}>{_rpr}'
                f'<w:t xml:space="preserve">Page </w:t></w:r>'
            )
            p.append(r)
            r = parse_xml(
                f'<w:r {nsdecls("w")}>{_rpr}'
                f'<w:fldChar w:fldCharType="begin"/></w:r>'
            )
            p.append(r)
            r = parse_xml(
                f'<w:r {nsdecls("w")}>{_rpr}'
                f'<w:instrText> PAGE </w:instrText></w:r>'
            )
            p.append(r)
            r = parse_xml(
                f'<w:r {nsdecls("w")}>{_rpr}'
                f'<w:fldChar w:fldCharType="end"/></w:r>'
            )
            p.append(r)
        elif copyright_text:
            # Copyright only
            r = parse_xml(
                f'<w:r {nsdecls("w")}>'
                f'<w:rPr><w:rFonts w:ascii="{FONT_BODY}" w:hAnsi="{FONT_BODY}"/>'
                f'<w:sz w:val="16"/><w:szCs w:val="16"/><w:color w:val="AAAAAA"/></w:rPr>'
                f'<w:t>{copyright_text}</w:t></w:r>'
            )
            p.append(r)

        footer._element.append(p)


# ---------------------------------------------------------------------------
# Template cover page helpers
# ---------------------------------------------------------------------------
def _template_has_cover_page(doc) -> bool:
    """Check if a template document has a cover page with a Title-styled paragraph."""
    for p in doc.paragraphs:
        if p.style and p.style.name == "Title":
            return True
    return False


def _fill_template_cover(doc, title: str, preamble_text: str = "",
                         date: str | None = None):
    """Fill the template's cover page placeholders and clear remaining body content.

    Looks for:
      - Paragraph with style "Title" -> replace text with title
      - First non-empty paragraph after Title -> replace with date (if provided)
      - Second non-empty paragraph after Title -> replace with preamble_text (if provided)

    Then removes all paragraphs and tables after the cover page content so the
    body can be appended fresh.
    """
    body = doc.element.body
    wns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    # Find the Title paragraph index in body children
    title_idx = None
    for i, child in enumerate(body):
        if child.tag == f"{wns}p":
            pPr = child.find(f"{wns}pPr")
            if pPr is not None:
                pStyle = pPr.find(f"{wns}pStyle")
                if pStyle is not None and pStyle.get(f"{wns}val") == "Title":
                    title_idx = i
                    break

    if title_idx is None:
        return  # no Title paragraph found

    # Replace Title text
    for p in doc.paragraphs:
        if p.style and p.style.name == "Title":
            # Clear existing runs and set new text
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = title
            else:
                p.add_run(title)
            break

    # Find non-empty paragraphs after Title to fill date and preamble
    post_title_filled = 0
    for i in range(title_idx + 1, len(body)):
        child = body[i]
        if child.tag != f"{wns}p":
            continue
        texts = child.findall(f".//{wns}t")
        txt = "".join(t.text or "" for t in texts).strip()
        if not txt:
            continue  # skip empty spacer paragraphs

        if post_title_filled == 0 and date:
            # First non-empty after Title -> date
            for t in texts:
                t.text = ""
            if texts:
                texts[0].text = date
            post_title_filled += 1
        elif post_title_filled <= 1 and preamble_text:
            # Second non-empty after Title -> preamble/validity
            for t in texts:
                t.text = ""
            if texts:
                texts[0].text = preamble_text
            post_title_filled += 1
        else:
            break

    # Determine last cover-page paragraph (last non-sectPr element before
    # content should be appended). We keep everything up to and including
    # the last existing paragraph/table.
    last_cover_idx = 0
    for i, child in enumerate(body):
        if child.tag == f"{wns}sectPr":
            continue
        last_cover_idx = i

    # Remove all body paragraphs and tables AFTER the cover page content,
    # but preserve the final sectPr.
    children = list(body)
    for child in children[last_cover_idx + 1:]:
        if child.tag == f"{wns}sectPr":
            continue  # keep final section properties
        body.remove(child)

    # Add a page break after cover page content
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    pb_para = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'<w:r><w:br w:type="page"/></w:r>'
        f'</w:p>'
    )
    # Insert before the final sectPr
    sect_pr = body.find(f"{wns}sectPr")
    if sect_pr is not None:
        sect_pr.addprevious(pb_para)
    else:
        body.append(pb_para)


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------
def convert(input_path: str, output_path: str, template_path: str | None = None,
            title: str | None = None, date: str | None = None,
            toc: bool = False, pagination: bool = True,
            copyright_text: str | None = None, skip_h1: bool = False,
            style_overrides: dict | None = None):
    input_path = os.path.abspath(input_path)
    output_path = os.path.abspath(output_path)
    input_dir = os.path.dirname(input_path)

    print(f"Reading: {input_path}")
    content = Path(input_path).read_text(encoding="utf-8")
    content = strip_front_matter(content)

    # --- Style resolution (priority: defaults < template < --style < inline < CLI) ---
    _apply_style({})  # reset to defaults

    # Extract style from template (first pass – reads the template independently)
    template_style: dict[str, str] = {}
    if template_path and os.path.isfile(template_path):
        template_style = extract_style_from_template(template_path)
        if template_style:
            print(f"  Template style extracted: {', '.join(template_style.keys())}")

    # Inline style hints from <!-- docx-style ... --> comment
    inline_style = parse_docx_style(content)
    content = strip_docx_style_comment(content)

    # Merge layers: template < --style file < inline < CLI flags
    merged_style: dict[str, str] = dict(template_style)
    if style_overrides:
        merged_style.update({k: v for k, v in style_overrides.items() if v is not None})
    merged_style.update(inline_style)
    if merged_style:
        _apply_style(merged_style)
        print(f"  Style applied: {', '.join(merged_style.keys())}")

    # --- Create document (second pass – uses template with correct styles) ---
    template_has_cover = False
    if template_path and os.path.isfile(template_path):
        print(f"Using template: {template_path}")
        doc = _open_template(template_path)
        # Check if template has a cover page (Title-styled paragraph)
        template_has_cover = _template_has_cover_page(doc)
        if template_has_cover:
            print("  Template cover page detected – preserving it")
        else:
            # No cover page – clear body as before (keep styles)
            for p in list(doc.paragraphs):
                p._element.getparent().remove(p._element)
            for t in list(doc.tables):
                t._element.getparent().remove(t._element)
    else:
        doc = Document()

    # Page margins – only override when there is NO template
    if not (template_path and os.path.isfile(template_path)):
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    with tempfile.TemporaryDirectory() as temp_dir:
        # Parse Markdown to AST using mistune
        md = mistune.create_markdown(
            renderer="ast",
            plugins=["table", "strikethrough"],
        )
        tokens: list[dict[str, Any]] = md(content)  # type: ignore[assignment]

        builder = DocxBuilder(doc, input_dir, temp_dir)

        # --- Title page logic ---
        # Find first H1 and first HR positions
        h1_idx = None
        h1_text = None
        first_hr_idx = None
        for i, tok in enumerate(tokens):
            if h1_idx is None and tok.get("type") == "heading" and tok.get("attrs", {}).get("level") == 1:
                h1_idx = i
                h1_text = builder._flatten_text(tok.get("children", []))
            if tok.get("type") == "thematic_break":
                first_hr_idx = i
                break

        # Determine effective title
        effective_title = title or h1_text  # --title overrides, else first H1

        # Extract preamble text for template cover (e.g. validity line)
        preamble_flat = ""
        if h1_idx is not None and first_hr_idx is not None:
            preamble_tokens = tokens[h1_idx + 1 : first_hr_idx]
            preamble_flat = "\n".join(
                builder._flatten_text(t.get("children", []))
                for t in preamble_tokens
                if t.get("type") == "paragraph"
            ).strip()

        if template_has_cover and effective_title:
            # Use the template's own cover page – fill placeholders
            _fill_template_cover(doc, effective_title,
                                 preamble_text=preamble_flat, date=date)
            print(f"  Template cover: {effective_title}" + (f" ({date})" if date else ""))

            # Body starts after the first HR (skip H1 + preamble + HR)
            if first_hr_idx is not None:
                tokens = tokens[first_hr_idx + 1:]
            elif h1_idx is not None:
                tokens = tokens[h1_idx + 1:]

        elif effective_title and h1_idx is not None and first_hr_idx is not None:
            # No template cover – generate our own title page
            preamble_tokens = tokens[h1_idx + 1 : first_hr_idx]

            _add_title_page(doc, builder, effective_title, preamble_tokens, date)
            print(f"  Title page: {effective_title}" + (f" ({date})" if date else ""))

            # Body starts after the first HR
            body_tokens = tokens[first_hr_idx + 1:]

            if not title:
                pass
            elif skip_h1:
                body_tokens = [t for t in body_tokens
                               if not (t.get("type") == "heading"
                                       and t.get("attrs", {}).get("level") == 1)
                               or t is not tokens[h1_idx]]
                pass
            else:
                body_tokens = [tokens[h1_idx]] + body_tokens

            tokens = body_tokens
        elif effective_title and h1_idx is not None:
            # H1 found but no HR – use H1 as title, rest is body
            _add_title_page(doc, builder, effective_title, [], date)
            print(f"  Title page: {effective_title}" + (f" ({date})" if date else ""))
            if title and not skip_h1:
                tokens = tokens
            else:
                tokens = tokens[h1_idx + 1:]

        # Table of Contents
        if toc:
            _add_toc(doc)
            print("  TOC: inserted (update in Word with Ctrl+A, F9)")

        # Build DOCX body
        builder.render_tokens(tokens)

        # Footer (pagination + copyright) – skip when template has its own
        if not template_has_cover and (pagination or copyright_text):
            _setup_footer(doc, pagination, copyright_text)

        # Save
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)

    print(f"Saved: {output_path}")
    print(f"  Mermaid diagrams processed: {builder.mermaid_counter}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX (no Pandoc required)")
    parser.add_argument("input", help="Input Markdown file")
    parser.add_argument("output", help="Output DOCX file")
    parser.add_argument("--template", help="Reference DOCX template for styles", default=None)
    parser.add_argument("--title", help="Title page text (adds a centred title page)", default=None)
    parser.add_argument("--date", help="Date shown on title page (requires --title)", default=None)
    parser.add_argument("--toc", action="store_true", help="Insert Table of Contents")
    parser.add_argument("--no-pagination", dest="pagination", action="store_false",
                        help="Disable page numbers in footer (enabled by default)")
    parser.add_argument("--copyright", help="Copyright text for footer", default=None)
    parser.add_argument("--skip-h1", dest="skip_h1", action="store_true",
                        help="Skip first H1 heading from body (use with --title)")
    # Style options
    parser.add_argument("--style", dest="style_file", default=None,
                        help="Style file (key: value format, same as inline comment)")
    parser.add_argument("--font-body", dest="font_body", default=None,
                        help="Body text font (default: Arial)")
    parser.add_argument("--font-heading", dest="font_heading", default=None,
                        help="Heading font (default: Arial)")
    parser.add_argument("--font-code", dest="font_code", default=None,
                        help="Code font (default: Consolas)")
    parser.add_argument("--font-size", dest="font_size", default=None,
                        help="Body text size in pt (default: 10.5)")
    parser.add_argument("--color-heading", dest="color_heading", default=None,
                        help="Heading color hex (default: 2D3B4D)")
    parser.add_argument("--color-body", dest="color_body", default=None,
                        help="Body text color hex (default: 333333)")
    parser.add_argument("--table-header-bg", dest="table_header_bg", default=None,
                        help="Table header background hex (default: D5E8F0)")
    parser.add_argument("--table-header-text", dest="table_header_text", default=None,
                        help="Table header text color hex (default: 2D3B4D)")
    parser.add_argument("--table-alt-row", dest="table_alt_row", default=None,
                        help="Table alternating row color hex (default: F2F2F2)")
    parser.add_argument("--table-border", dest="table_border", default=None,
                        help="Table border color hex (default: CCCCCC)")
    parser.add_argument("--table-border-size", dest="table_border_size", default=None,
                        help="Table border width in half-points (default: 4)")
    parser.add_argument("--table-cell-margin", dest="table_cell_margin", default=None,
                        help="Table cell margin in twips (default: 28)")
    parser.add_argument("--table-font-size", dest="table_font_size", default=None,
                        help="Table text size in pt (default: 9.5)")
    parser.add_argument("--no-banded-rows", dest="table_banded_rows",
                        action="store_const", const="false", default=None,
                        help="Disable alternating row shading")
    parser.add_argument("--code-bg", dest="code_bg", default=None,
                        help="Code block background hex (default: F5F5F5)")
    parser.add_argument("--code-font-size", dest="code_font_size", default=None,
                        help="Code text size in pt (default: 9)")
    args = parser.parse_args()

    if not os.path.isfile(args.input):
        print(f"Error: Input file not found: {args.input}")
        sys.exit(1)

    # Build style overrides: --style file < CLI flags
    style_overrides = {}
    if args.style_file and os.path.isfile(args.style_file):
        style_overrides.update(parse_style_file(args.style_file))
    style_keys = [
        "font_body", "font_heading", "font_code", "font_size",
        "color_heading", "color_body",
        "table_header_bg", "table_header_text", "table_alt_row",
        "table_border", "table_border_size", "table_cell_margin",
        "table_font_size", "table_banded_rows",
        "code_bg", "code_font_size",
    ]
    for key in style_keys:
        val = getattr(args, key, None)
        if val is not None:
            style_overrides[key] = val

    convert(args.input, args.output, args.template,
            title=args.title, date=args.date,
            toc=args.toc, pagination=args.pagination,
            copyright_text=args.copyright, skip_h1=args.skip_h1,
            style_overrides=style_overrides if style_overrides else None)


if __name__ == "__main__":
    main()
